#!/usr/bin/env python3
"""Generate ATS-safe DOCX resume for Aniket Kumar — Senior DevOps Engineer."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/crypticani/projects/crypticani/crypticani.github.io/public/files/Aniket_Kumar_DevOps_Resume.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = Pt(12.5)

BLUE = RGBColor(0x1b, 0x3a, 0x5c)
GRAY = RGBColor(0x44, 0x44, 0x44)
GRAY2 = RGBColor(0x55, 0x55, 0x55)
GRAY3 = RGBColor(0x66, 0x66, 0x66)


def add_heading_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1B3A5C')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body(doc, text, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12.5)
    run = p.add_run(text)
    run.font.size = size
    run.font.name = 'Calibri'


def add_bullet(doc, text, size=Pt(10)):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = size
    run.font.name = 'Calibri'


def add_skill_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = Pt(12)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = 'Calibri'
    r2 = p.add_run(" " + value)
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'


def add_job_header(doc, title, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = 'Calibri'
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.4), alignment=2)
    p.add_run('\t')
    r2 = p.add_run(dates)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GRAY
    r2.font.name = 'Calibri'


# ========== HEADER ==========
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(1)
run = name.add_run("Aniket Kumar")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(1)
run = subtitle.add_run("Senior DevOps Engineer  |  Platform Engineering  |  Cloud Infrastructure  |  SRE")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run.font.name = 'Calibri'

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(3)
run = contact.add_run(
    "Samastipur, Bihar, India  \u2022  +91 9939093867  \u2022  aniket.kumar@hotmail.com  \u2022  "
    "linkedin.com/in/crypticani  \u2022  github.com/crypticani  \u2022  crypticani.dev"
)
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

# ========== PROFESSIONAL SUMMARY ==========
add_heading_line(doc, "Professional Summary")
add_body(doc,
    "Senior DevOps Engineer and acting technical lead with 4+ years of experience in high-availability "
    "infrastructure, CI/CD automation, Kubernetes orchestration, and production operations across multi-cloud "
    "and on-premise environments. Lead a team of 8+, driving incident management, root cause analysis, cloud "
    "migration strategies, and DevSecOps practices while defining SLOs and sustaining 99.9% SLA across 15+ production applications. "
    "Full-stack development background with strong end-to-end troubleshooting skills; expertise in Infrastructure "
    "as Code, observability, database HA/DR, cost optimization, and security compliance (ISO 27001, SOC 2).")

# ========== TECHNICAL SKILLS ==========
add_heading_line(doc, "Technical Skills")
add_skill_line(doc, "Cloud & Infrastructure:", "AWS, OCI, Azure, GCP, Cloud Migration, Disaster Recovery, Cost Optimization, High Availability")
add_skill_line(doc, "Containers, CI/CD & IaC:", "Docker, Kubernetes, Docker Compose, Jenkins, GitHub Actions, GitLab CI/CD, Terraform, Ansible, Configuration Management, Grype")
add_skill_line(doc, "Observability & Databases:", "Prometheus, Grafana, Loki, Alertmanager, PromQL, PostgreSQL, Patroni, MySQL, MongoDB, Redis, KeyDB, ClickHouse")
add_skill_line(doc, "Security, Compliance & Systems:", "Keycloak, SAML/OIDC, SSL/TLS, RBAC, DevSecOps, ISO 27001, SOC 2, Linux, Nginx, HAProxy, DNS, Python, Bash, Git")
add_skill_line(doc, "Process & ITSM:", "Agile/Scrum, Jira, Freshservice, ITSM, Change Management, Incident Management, Release Management")

# ========== PROFESSIONAL EXPERIENCE ==========
add_heading_line(doc, "Professional Experience")

# --- Current Role ---
add_job_header(doc, "DevOps Engineer (Acting Technical Lead) \u2014 Kochar Innovations Pvt Ltd", "Mar 2025 \u2013 Present")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("Amritsar, India")
r.font.size = Pt(9.5)
r.font.color.rgb = GRAY2
r.font.name = 'Calibri'

add_bullet(doc, "Serve as acting technical lead for 15+ production applications, mentoring 5+ engineers, leading root cause analysis, coordinating incident response via Freshservice, and driving Agile sprint planning while defining SLOs and sustaining 99.9% uptime SLA.")
add_bullet(doc, "Planned and executed cloud-to-cloud and cloud-to-on-premise migrations with near-zero-downtime cutover strategies, post-migration validation, and stabilization workflows.")
add_bullet(doc, "Drove cloud cost optimization initiatives through rightsizing, idle resource cleanup, and reserved capacity planning, reducing monthly infrastructure spend by 25%.")
add_bullet(doc, "Supported ISO 27001 and SOC 2 audit processes by enforcing access controls, change management workflows, centralized logging, backup validation, and security policy compliance across cloud infrastructure.")

# --- Previous Role ---
add_job_header(doc, "DevOps Engineer \u2014 Kochar Infotech Ltd", "Jun 2023 \u2013 Feb 2025")
note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(1)
r = note.add_run("Continuous role; company restructured to Kochar Innovations Pvt Ltd with same team and responsibilities.")
r.font.size = Pt(9)
r.font.color.rgb = GRAY3
r.font.name = 'Calibri'
r.italic = True

add_bullet(doc, "Designed and maintained CI/CD pipelines using Jenkins, GitHub Actions, GitLab CI/CD, Terraform, and Ansible, increasing release cadence from bi-weekly to multiple daily deployments.")
add_bullet(doc, "Containerized 20+ services with multi-stage Dockerfiles and managed Kubernetes workloads with rolling updates and zero-downtime releases; reduced image size by 60% and deployment time by 50%.")
add_bullet(doc, "Integrated Grype-based container vulnerability scanning into CI/CD security gates, blocking critical/high CVE deployments and reducing production vulnerability exposure by 90%.")
add_bullet(doc, "Built full-stack observability using Prometheus, Grafana, Alertmanager, Loki, and custom Python exporters across 100+ hosts and 200+ containers with 15+ operational dashboards.")
add_bullet(doc, "Architected HA data platforms: PostgreSQL + Patroni, MySQL Group Replication, MongoDB Replica Sets, ClickHouse, Redis Sentinel, and KeyDB with automated backup, PITR, and failover runbooks.")
add_bullet(doc, "Implemented disaster recovery with automated snapshots, cross-region backup replication, and documented recovery runbooks, achieving sub-30-minute RTO.")

# --- Intern ---
add_job_header(doc, "IT Intern \u2014 Kochar Infotech Ltd", "Jun 2022 \u2013 May 2023")
add_bullet(doc, "Built organization-wide monitoring infrastructure from scratch using Prometheus, Grafana, Loki, and Fluent Bit; automated backup workflows to AWS S3; developed Django + MQTT device portal for 50+ devices.")

# ========== CERTIFICATIONS & EDUCATION ==========
add_heading_line(doc, "Certifications & Education")

p1 = doc.add_paragraph()
p1.paragraph_format.space_after = Pt(1)
p1.paragraph_format.line_spacing = Pt(12)
r = p1.add_run("Certifications: ")
r.bold = True
r.font.size = Pt(9.5)
r.font.name = 'Calibri'
r2 = p1.add_run("Oracle Cloud Infrastructure 2025 Certified Architect Associate \u2022 Cloud Architecture \u2013 Google Cloud Skills Boost \u2022 Google IT Automation with Python \u2013 Coursera \u2022 Foundations of Project Management \u2013 Coursera")
r2.font.size = Pt(9.5)
r2.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(0)
p2.paragraph_format.line_spacing = Pt(12)
r = p2.add_run("Education: ")
r.bold = True
r.font.size = Pt(9.5)
r.font.name = 'Calibri'
r2 = p2.add_run("MCA, Lovely Professional University, Punjab (2021\u20132023), CGPA 8.97 \u2022 BCA, Dev Sanskriti Vishwavidyalaya, Haridwar (2018\u20132021), CGPA 8.0, Academic Excellence Award")
r2.font.size = Pt(9.5)
r2.font.name = 'Calibri'

doc.save(OUTPUT)
print(f"DOCX saved to: {OUTPUT}")
